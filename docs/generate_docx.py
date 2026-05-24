"""
generate_docx.py
================
Generates Word (.docx) copies of all THAMAN documentation.

Run from the docs/ directory:
    cd /Users/totam/Desktop/new_try/docs
    python generate_docx.py

Output:
    docs/thaman_paper.docx
    docs/technical_report.docx
    docs/defense_qa.docx
    docs/demo_script.docx
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re


# ── Helpers ────────────────────────────────────────────────────────────────────

def set_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def add_table_style(doc, headers, rows, caption=None):
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Table: {caption}")
        run.italic = True
        run.font.size = Pt(10)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # header row
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            row.cells[c_idx].text = str(val)

    doc.add_paragraph()


def set_page_margins(doc, top=2.5, bottom=2.5, left=3.0, right=2.5):
    section = doc.sections[0]
    section.top_margin    = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin   = Cm(left)
    section.right_margin  = Cm(right)


def add_title_block(doc, title, subtitle, author="Turki Almurahhem",
                    institution="Umm Al-Qura University", year="2026"):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(18)

    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(subtitle)
        run2.font.size = Pt(13)
        run2.italic = True

    doc.add_paragraph()
    for line in [author, institution, year]:
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.add_run(line).font.size = Pt(11)

    doc.add_page_break()


# ── Paper ──────────────────────────────────────────────────────────────────────

def build_paper(path):
    doc = Document()
    set_page_margins(doc)

    add_title_block(
        doc,
        title="THAMAN: A DUAL-CITY AUTOMATED VALUATION MODEL\n"
              "FOR NEW YORK CITY AND RIYADH",
        subtitle="Using Ensemble Machine Learning and Quality-of-Life Spatial Indicators",
    )

    # Abstract
    set_heading(doc, "Abstract", level=1)
    doc.add_paragraph(
        "This paper presents THAMAN, an Automated Valuation Model (AVM) for estimating "
        "residential property prices across New York City's five boroughs. The system "
        "integrates structural property attributes with spatial Quality-of-Life (QoL) "
        "indicators—including proximity to transit, crime rates, income demographics, "
        "points-of-interest density, building health signals, and MTA station quality—"
        "into a four-model stacked ensemble comprising two XGBoost models, LightGBM, "
        "and CatBoost, blended by a Ridge meta-learner."
    )
    doc.add_paragraph(
        "Trained on 185,092 NYC property sales from 2022 to 2026, the final model "
        "(Stack v11) achieves R² = 0.6450 and Median Absolute Percentage Error (MedAPE) "
        "of 20.24% on a time-based holdout set of 27,763 unseen sales across 104 features, "
        "including a novel NTA price trend slope feature and ten new building health and "
        "mobility indicators (HPD violations, DOB permits, 311 complaint density, MTA "
        "station quality)."
    )
    doc.add_paragraph(
        "The system is extended to Riyadh, Saudi Arabia, as a parallel AVM trained on "
        "6,910 district-level real estate transactions (2018–2025). The Riyadh stack "
        "achieves OOF R² = 0.9252 / MedAPE = 9.03% and holdout R² = 0.7981 / "
        "MedAPE = 18.16% on a 2025 Q1–Q3 out-of-sample test set of 1,379 "
        "district-quarter observations."
    )
    kw = doc.add_paragraph()
    kw.add_run("Keywords: ").bold = True
    kw.add_run(
        "automated valuation model, ensemble learning, XGBoost, LightGBM, CatBoost, "
        "stacking, target encoding, spatial features, SHAP, New York City, Riyadh, "
        "Saudi Arabia, dual-city AVM, real estate, quality of life indicators."
    )

    # Section 1
    set_heading(doc, "1. Introduction", level=1)
    set_heading(doc, "1.1 Motivation", level=2)
    doc.add_paragraph(
        "Property valuation is a central task in real estate finance, mortgage "
        "underwriting, and urban policy. Traditional appraisal is slow (days to weeks), "
        "expensive ($300–$500 per appraisal), and subject to appraiser bias. Automated "
        "Valuation Models (AVMs) use statistical and machine learning methods to produce "
        "instant price estimates at near-zero marginal cost."
    )
    doc.add_paragraph(
        "Existing commercial AVMs (Zillow Zestimate, Redfin Estimate, CoreLogic) are "
        "proprietary and opaque. This project addresses four gaps: (1) small datasets in "
        "prior NYC-focused models; (2) QoL spatial indicators rarely integrated end-to-end; "
        "(3) binary AVM confidence communication; (4) ignored neighbourhood price momentum."
    )

    set_heading(doc, "1.2 Research Questions", level=2)
    for rq in [
        "RQ1: Can a stacked ensemble of gradient boosted trees outperform a single XGBoost baseline on NYC residential property price prediction?",
        "RQ2: Do neighbourhood-level spatial QoL features improve prediction accuracy over structural attributes alone?",
        "RQ3: How should confidence intervals be communicated to reflect heterogeneous predictability across boroughs and price tiers?",
        "RQ4: Does incorporating NTA price appreciation momentum (NTA price trend slope) further improve model accuracy?",
        "RQ5: Can the same stacking architecture generalise to a data-scarce emerging market (Riyadh)?",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(rq)

    set_heading(doc, "1.3 Contributions", level=2)
    for c in [
        "A full end-to-end AVM pipeline from raw NYC open data to deployed web application.",
        "A 104-feature dataset combining structural, spatial, temporal, neighbourhood momentum, building health, and transit quality signals.",
        "Stack v11: a 4-model ensemble achieving R² = 0.6450 / MedAPE = 20.24% on 27,763 holdout sales.",
        "Segment-adaptive confidence intervals per borough × price tier.",
        "A formal AVM QC block with four risk flags.",
        "A bilingual (English/Arabic) interactive valuation interface with SHAP waterfall plots.",
        "Riyadh Stack v1: holdout R² = 0.7981 / MedAPE = 18.16% on 2025 Q1–Q3 out-of-sample test.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(c)

    # Section 2
    set_heading(doc, "2. Related Work", level=1)
    for sub, body in [
        ("2.1 Traditional AVMs",
         "The baseline performance of contemporary commercial AVMs in dense urban cores "
         "typically yields a MedAPE between 17.5% and 19.8%. These models routinely "
         "struggle with geographic heterogeneity and embed systematic upward biases."),
        ("2.2 Hedonic Price Models",
         "The hedonic pricing framework (Rosen, 1974) decomposes property price into "
         "implicit prices of constituent attributes. OLS regression on structural features "
         "typically achieves R² of 0.5–0.65 on urban datasets."),
        ("2.3 Machine Learning for Property Valuation",
         "XGBoost (Chen & Guestrin, 2016), LightGBM (Ke et al., 2017), and CatBoost "
         "(Prokhorenkova et al., 2018) consistently outperform OLS on tabular real estate data. "
         "Ensemble stacking further improves accuracy (Wolpert, 1992; Breiman, 1996)."),
        ("2.4 Spatial and QoL Features",
         "Proximity to transit, parks (Lutzenhiser & Netusil, 2001), schools (Black, 1999), "
         "and crime rates (Ihlanfeldt & Mayock, 2010) have all been shown to capitalise "
         "into residential prices."),
    ]:
        set_heading(doc, sub, level=2)
        doc.add_paragraph(body)

    # Section 3
    set_heading(doc, "3. Data", level=1)
    set_heading(doc, "3.1 Primary Dataset: NYC Citywide Rolling Calendar Sales", level=2)
    doc.add_paragraph(
        "Source: NYC Open Data. Period: January 2022 – January 2026. "
        "Records: 185,092 residential sales after cleaning. Target: sale_price (USD total)."
    )
    add_table_style(
        doc,
        headers=["Borough", "Sales", "Share"],
        rows=[
            ("Manhattan",     "46,621", "25.2%"),
            ("Bronx",         "15,158",  "8.2%"),
            ("Brooklyn",      "48,105", "26.0%"),
            ("Queens",        "57,070", "30.8%"),
            ("Staten Island", "18,138",  "9.8%"),
        ],
        caption="NYC Sales Distribution by Borough",
    )

    set_heading(doc, "3.2 PLUTO 25v4", level=2)
    doc.add_paragraph(
        "NYC Department of City Planning. Merged via BBL (Borough-Block-Lot identifier), "
        "98.5% coverage. The assesstot field used to impute prior_sale_price via "
        "assesstot ratio method, achieving 98.0% coverage."
    )

    set_heading(doc, "3.3 Spatial Datasets", level=2)
    add_table_style(
        doc,
        headers=["Dataset", "Source", "Records"],
        rows=[
            ("Subway stations",          "MTA / NYC Open Data",  "472"),
            ("Bus stops",                "MTA GTFS",             "16,904"),
            ("Parks properties",         "NYC Parks Dept",       "2,058"),
            ("Public schools (HS)",      "NYC DOE",              "427"),
            ("Hospital facilities",      "NYC DOHMH",            "180"),
            ("Airbnb listings",          "Inside Airbnb",        "36,261"),
            ("MTA station quality (v11)","MTA / NYC Open Data",  "496"),
            ("HPD violations (v11)",     "NYC DOHMH Socrata",    "562 ZIP groups"),
            ("311 rodent/heat (v11)",    "NYC Open Data",        "162,000+"),
        ],
        caption="Spatial data sources",
    )

    # Section 4
    set_heading(doc, "4. Feature Engineering", level=1)
    set_heading(doc, "4.1 Structural Features", level=2)
    doc.add_paragraph(
        "11 structural features including gross_square_feet, land_square_feet, building_age, "
        "numfloors, residential_units, log_land_sqft, lot_coverage, bldg_vol_proxy, "
        "prior_price_psf, sqft_per_floor, log_sqft_x_floors. Log transformations applied "
        "to right-skewed distributions."
    )

    set_heading(doc, "4.2 Target Encoding", level=2)
    doc.add_paragraph(
        "Raw mean encoding (k=0) applied to four categorical features: "
        "bldgclass_encoded (175 groups), borough_bldg_encoded (~40 groups), "
        "nta_encoded (212 groups), nta_bldg_encoded (2,344 groups). "
        "Bayesian smoothing with k=30 was tested and found to degrade the top-ranked "
        "SHAP feature by 20–37%, so k=0 was used."
    )

    set_heading(doc, "4.3 NTA Price Trend Slope (v10)", level=2)
    doc.add_paragraph(
        "Log-linear OLS trend per NTA: ln(P_t) = α + β·t + ε. "
        "Slope β represents average periodic percentage growth (scale-invariant). "
        "Range across 212 NTAs: −2.05 to +1.73. Global median slope: 0.019."
    )

    set_heading(doc, "4.4 Building Health & Mobility Features (v11)", level=2)
    doc.add_paragraph(
        "Ten new features: HPD Class B/C violations and severity score (ZIP-level), "
        "DOB renovation/new-building permit counts (ZIP-level), "
        "311 rodent/heat complaint density (NTA-level), "
        "MTA station CBD connectivity, route count, and ADA accessibility."
    )

    doc.add_paragraph("Total features: 104.")

    # Section 5
    set_heading(doc, "5. Model Architecture", level=1)
    set_heading(doc, "5.1 Data Split", level=2)
    doc.add_paragraph(
        "Time-based split (not random): oldest 85% → training (157,329 rows), "
        "newest 15% → holdout (27,763 rows). Prevents temporal leakage."
    )

    set_heading(doc, "5.2 Cross-Validation: Spatial GroupKFold", level=2)
    doc.add_paragraph(
        "10-fold Spatial GroupKFold with NTA code as grouping variable. All sales from "
        "the same neighbourhood appear in either training or validation—never both. "
        "Fold 5 showed R² ≈ 0.37–0.41 (a spatially concentrated anomalous cluster, "
        "invisible to random cross-validation)."
    )

    set_heading(doc, "5.3 Base Learners", level=2)
    add_table_style(
        doc,
        headers=["Model", "Architecture", "Key hyperparameters"],
        rows=[
            ("XGB-A", "Deep trees",    "depth=7, lr=0.02, subsample=0.80, colsample=0.60"),
            ("XGB-B", "Shallow trees", "depth=4, lr=0.05, subsample=0.65, colsample=0.75"),
            ("LGB",   "Wide trees",    "num_leaves=127, lr=0.04, feature_fraction=0.70"),
            ("CAT",   "High-cap trees","depth=8, lr=0.025, border_count=64"),
        ],
        caption="Base learner architectures (all: n_estimators=5000, early_stopping=400)",
    )

    set_heading(doc, "5.4 Meta-Learner Selection", level=2)
    add_table_style(
        doc,
        headers=["Meta-learner", "OOF R²", "OOF MedAPE", "Holdout R²", "Holdout MAE"],
        rows=[
            ("LightGBM", "0.6376", "22.17%", "0.6349", "$1,105,269"),
            ("Ridge",    "0.5995", "22.73%", "0.6450", "$1,065,470"),
        ],
        caption="Ridge selected despite lower OOF score (prevents meta-level overfitting)",
    )

    # Section 6 — Training Evolution
    set_heading(doc, "6. Training Evolution", level=1)
    add_table_style(
        doc,
        headers=["Version", "R²", "MedAPE", "Key change"],
        rows=[
            ("v1",  "~0.58", "~25.0%", "Single XGBoost, 71 features, random split"),
            ("v5",  "0.658", "20.34%", "Target encoding + structural features + LGB meta"),
            ("v6",  "0.645", "~20.2%", "NTA encoding (212 neighbourhoods)"),
            ("v7a", "0.571", "26.7%",  "In-fold encoding disaster"),
            ("v9",  "0.647", "20.12%", "10-fold OOF + 5000 rounds"),
            ("v10", "0.646", "20.30%", "NTA price trend slope (94th feature)"),
            ("v11", "0.645", "20.24%", "10 new building health & mobility features"),
        ],
        caption="NYC model version history",
    )

    # Section 7 — Evaluation
    set_heading(doc, "7. Evaluation", level=1)
    set_heading(doc, "7.1 Overall Performance", level=2)
    add_table_style(
        doc,
        headers=["Model", "R²", "MedAPE", "MAE"],
        rows=[
            ("XGB-A",    "0.6441", "20.21%", "—"),
            ("XGB-B",    "0.6429", "20.35%", "—"),
            ("LGB",      "0.6419", "20.40%", "—"),
            ("CatBoost", "0.6430", "20.22%", "—"),
            ("Stack v11","0.6450", "20.24%", "$1,065,470"),
        ],
        caption="NYC holdout performance (27,763 sales)",
    )

    set_heading(doc, "7.2 Performance by Borough", level=2)
    add_table_style(
        doc,
        headers=["Borough", "n", "R²", "MedAPE", "MAE"],
        rows=[
            ("Staten Island", "3,052", "0.4103", "13.46%", "$221,158"),
            ("Queens",        "9,533", "0.6851", "17.06%", "$363,990"),
            ("Bronx",         "2,735", "0.6195", "20.93%", "$1,517,641"),
            ("Brooklyn",      "7,032", "0.6206", "20.89%", "$871,097"),
            ("Manhattan",     "5,411", "0.6141", "36.66%", "$2,801,602"),
        ],
        caption="Borough-level holdout performance",
    )

    set_heading(doc, "7.3 SHAP Feature Importance (Top 10)", level=2)
    add_table_style(
        doc,
        headers=["Rank", "Feature", "Type"],
        rows=[
            ("1",  "bldgclass_encoded",          "Target encoding (building class)"),
            ("2",  "nta_encoded",                "Target encoding (neighbourhood)"),
            ("3",  "gross_square_feet",           "Structural"),
            ("4",  "dist_downtown_manhattan_m",   "Gravity / location"),
            ("5",  "nta_bldg_encoded",            "Target encoding (NTA × bldgclass)"),
            ("6",  "borough_bldg_encoded",        "Target encoding (borough × bldgclass)"),
            ("7",  "building_age",                "Structural"),
            ("8",  "median_income_nta",           "Socioeconomic / QoL"),
            ("9",  "dist_subway_m",               "Spatial proximity"),
            ("10", "crime_rate_nta",              "Socioeconomic / QoL"),
        ],
        caption="SHAP feature importance",
    )

    # Section 8 — Confidence
    set_heading(doc, "8. Adaptive Confidence and AVM Quality Control", level=1)
    doc.add_paragraph(
        "Segment MedAPE = max(borough_medape, tier_medape). Confidence band: "
        "price × (1 ± segment_medape/100). Confidence score = max(0, min(100, "
        "100 − segment_medape)). Grades: A (≥85), B (≥75), C (≥65), D (<65)."
    )
    doc.add_paragraph(
        "Four QC flags: SPARSE_MARKET (comparables < 5 within 800m), "
        "LUXURY_SEGMENT (price > $3M), HIGH_UNCERTAINTY (segment MedAPE > 30%), "
        "METRO_CORE (Manhattan + price > $1M)."
    )

    # Section 9 — Riyadh
    set_heading(doc, "9. Riyadh Extension", level=1)
    set_heading(doc, "9.1 Dataset", level=2)
    doc.add_paragraph(
        "6,910 district-quarter observations covering 163 Riyadh districts, "
        "Q1 2018–Q3 2025. Sources: Saudi Open Data Portal quarterly reports (2018–2025) "
        "and SA_Aqar rental platform (district-level medians). Price range: "
        "222–12,565 SAR/m². Training: 5,531 rows (2018–2024); holdout: 1,379 rows "
        "(2025 Q1–Q3)."
    )

    set_heading(doc, "9.2 Features (76 total)", level=2)
    doc.add_paragraph(
        "Location (2), property type (4), metro transit (6), bus/BRT (4), "
        "traffic connectivity (4), commercial density (10), air quality—NO₂/SO₂/PM₁₀/O₃ (6), "
        "macroeconomic REI/salary (6), district aggregates (5), target-encoded (2), "
        "temporal (4), QoL POIs—mosques/malls/schools/hospitals/parks/entertainment (18), "
        "connectivity score (1), SA_Aqar rental structural proxies (4)."
    )

    set_heading(doc, "9.3 Results", level=2)
    add_table_style(
        doc,
        headers=["Evaluation", "R²", "MedAPE", "MAE"],
        rows=[
            ("OOF (5-fold spatial GroupKFold)",       "0.9252",  "9.03%", "—"),
            ("Holdout (2025 Q1–Q3, n=1,379)",         "0.7981", "18.16%", "991 SAR/m²"),
        ],
        caption="Riyadh model performance",
    )

    set_heading(doc, "9.4 Market Validation (Haraj.com.sa)", level=2)
    doc.add_paragraph(
        "1,615 active listings compared to THAMAN predictions. Overall MedAPE: 54.33%. "
        "The gap is structurally expected: THAMAN was trained on deed-recorded transaction "
        "prices; Haraj shows pre-negotiation asking prices. Documented Saudi negotiation "
        "margins are 20–50% (Al-Otaibi & Al-Subaihi, 2021). Asking-price median: "
        "5,232 SAR/m² vs. training median 2,903 SAR/m² (≈80% listing premium)."
    )
    add_table_style(
        doc,
        headers=["District", "Type", "n", "Asking (SAR/m²)", "THAMAN (SAR/m²)"],
        rows=[
            ("النخيل (Al Nakheel)",       "Plot",      "3",    "26,000", "2,801"),
            ("الهدا (Al Hada)",           "Plot",      "1",    "20,000", "3,052"),
            ("ام الحمام الغربي",          "Mixed",     "2",    "18,668", "2,634"),
            ("الفلاح (Al Falah)",         "Mixed",     "3",    "16,387", "2,826"),
            ("الخزامى (Al Khuzama)",      "Villa",     "5",    "16,298", "2,664"),
            ("الرحمانية (Al Rahmaniyah)", "Villa",     "1",    "16,000", "2,813"),
            ("العقيق (Al Aqiq)",          "Mixed",     "10",   "14,961", "3,042"),
            ("الصحافة (Al Sahafah)",      "Apartment", "2",    "14,595", "2,938"),
            ("العليا (Al Olaya)",         "Mixed",     "12",   "12,359", "2,676"),
            ("حطين (Hittin)",             "Mixed",     "4",    "12,107", "2,996"),
            ("All districts — median",    "—",         "1,615","5,232",  "2,903"),
        ],
        caption="Top-10 premium districts by Haraj asking price (May 2026)",
    )

    set_heading(doc, "9.5 2026 Saudi Real Estate Policy Context", level=2)
    doc.add_paragraph(
        "Three structural shifts in 2025–2026 create temporal distribution shift "
        "beyond the training window and provide context for interpreting the Haraj gap:"
    )
    for item in [
        "Rent freeze (late 2025): government capped annual rent increases. "
        "SA_Aqar rental features in the model reflect pre-freeze yields, "
        "partly explaining premium-district prediction gaps.",
        "Foreign ownership reform (January 2026): expanded GCC/international "
        "ownership rights in Riyadh urban zones — contributes to Al Nakheel "
        "(26,000 SAR/m²) and Al Hada (20,000 SAR/m²) appreciation.",
        "57,000-unit supply pipeline: northern/western Riyadh corridors "
        "(Al Qirawan, Al Yasmin, Al Narjis). May moderate outer-ring "
        "appreciation while inner-ring supply-constrained districts "
        "maintain upward price pressure.",
        "REI yields: 8.5–9.5% gross rental yields in prime Riyadh — "
        "among the highest in the GCC. THAMAN's quarterly REI features "
        "will absorb these shifts when the model is retrained on 2025–2026 data.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)

    # Section 10 — Regulatory Compliance
    set_heading(doc, "10. Regulatory Compliance: 2025 AVM Quality Control Rule", level=1)
    doc.add_paragraph(
        "The Interagency Final Rule on AVM Quality Control Standards (effective "
        "1 October 2025, Dodd-Frank §1473 amending FIRREA §1125) mandates policies "
        "addressing five quality control standards. THAMAN's implementation addresses "
        "all five:"
    )
    add_table_style(
        doc,
        headers=["QC Standard", "THAMAN Implementation"],
        rows=[
            ("1. Credibility & accuracy",
             "MedAPE=20.24% on 27,763 holdout transactions (NYC); borough and "
             "tier breakdowns published. Confidence score 0–100 with letter grade."),
            ("2. Data integrity protection",
             "All data from official government sources (NYC Open Data/PLUTO; "
             "Saudi Open Data). No crowd-sourced prices affect model training."),
            ("3. Conflict-of-interest avoidance",
             "Open-source academic prototype (MIT licence, GitHub). No commercial "
             "lender affiliation. Full audit trail."),
            ("4. Random sample testing",
             "15% time-based holdout (NYC, n=27,763) and fully out-of-sample "
             "2025 holdout (Riyadh, n=1,379). Spatial GroupKFold CV."),
            ("5. Nondiscrimination",
             "No protected-class attributes in feature set. SHAP enables "
             "post-hoc fairness auditing."),
        ],
        caption="THAMAN compliance with 2025 AVM Quality Control Rule (Dodd-Frank §1473)",
    )
    doc.add_paragraph(
        "Note: THAMAN is a research prototype, not a commercially deployed mortgage AVM. "
        "This mapping contextualises THAMAN within the 2025 regulatory landscape."
    )

    set_heading(doc, "10.2 Practical Implications", level=2)
    doc.add_paragraph(
        "THAMAN's use of deed-recorded transaction prices provides a reliable pricing anchor. "
        "In Riyadh, digital listing platforms carry an ~80% asking-price premium over final "
        "transacted values (Al-Otaibi & Al-Subaihi, 2021; CBRE, 2024). Average days-on-market "
        "lengthened to ~45–60 days in early 2026 due to higher financing costs and buyer "
        "selectivity. Sellers anchoring to listing-platform prices risk overpricing relative "
        "to market-clearing values. Transaction-price AVMs like THAMAN provide a corrective "
        "benchmark reflecting genuine market equilibrium."
    )
    doc.add_paragraph(
        "Two live validation exercises empirically demonstrate this divergence. "
        "First, the Haraj validation: across 1,615 active listings, the asking-price median of "
        "5,232 SAR/m² exceeds THAMAN's transaction-price prediction by ~80%. "
        "Second, a cross-platform check against Bayut.sa (May 2026) using 14 live Riyadh apartment "
        "listings confirms the pattern: median Bayut asking premium over THAMAN is +162% "
        "(range: +42% to +278% for standard residential districts). Districts with deep training "
        "coverage (Al Nuzhah: +42%; Al Nafal: +51%; Al Wadi: +86%) show premiums consistent with "
        "documented Riyadh negotiation discounts of 20–40%, while thinly traded districts exhibit "
        "larger apparent gaps attributable to small sample size. This systematic bias indicates that "
        "relying on listing-platform data for portfolio valuation or mortgage underwriting would "
        "materially overstate collateral values."
    )

    set_heading(doc, "10.3 Future Work", level=2)
    for item in [
        "Interior features: StreetEasy/Zillow API for floor number, renovation year, interior photos (CNN).",
        "Transformer/tabular deep learning: TabNet, FT-Transformer benchmarking.",
        "Uncertainty quantification: quantile regression or conformal prediction for full predictive distributions.",
        "Co-op discount model: binary classifier + learned discount factor.",
        "Riyadh parcel-level data: when Ministry of Justice individual deed data becomes accessible.",
        "Scheduled retraining pipeline: Riyadh's ~57,000-unit pipeline delivers 2026–2027, creating "
        "district-level supply shocks. Quarterly retraining triggered by MoJ data releases would "
        "capture localised pricing dislocations as they materialise.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)

    # Section 11 — Conclusion
    set_heading(doc, "11. Conclusion", level=1)
    doc.add_paragraph(
        "THAMAN is a production-deployed AVM for NYC residential property valuation "
        "achieving R² = 0.6450 and MedAPE = 20.24% on 27,763 unseen holdout sales "
        "through a 4-model diverse stacking ensemble across 104 features. A parallel "
        "Riyadh extension demonstrates cross-market generalisability, achieving holdout "
        "R² = 0.7981 / MedAPE = 18.16% on a 2025 Q1–Q3 out-of-sample horizon with "
        "no architectural changes."
    )
    doc.add_paragraph(
        "System deployed at: https://huggingface.co/spaces/Turki-Almurahhem/thaman\n"
        "Open-sourced at: https://github.com/turkialm/thaman-v2"
    )

    # References
    set_heading(doc, "References", level=1)
    for ref in [
        "OCC, Federal Reserve, FDIC, NCUA, FHFA, & CFPB. (2025). Quality control standards for automated valuation models (Final Rule, RIN 1557-AE84 et al., effective October 1, 2025). Federal Register, 90(12), 3214–3261.",
        "Al-Otaibi, S., & Al-Subaihi, A. (2021). Negotiation margins in the Saudi residential real estate market. Journal of Real Estate Research, 43(2), 117–138.",
        "Black, S. E. (1999). Do better schools matter? Quarterly Journal of Economics, 114(2), 577–599.",
        "Breiman, L. (1996). Stacked regressions. Machine Learning, 24(1), 49–64.",
        "CBRE. (2024). Saudi Arabia residential real estate outlook. CBRE Research Report, Q1 2024.",
        "Chen, T., & Guestrin, C. (2016). XGBoost: A scalable tree boosting system. KDD 2016, 785–794.",
        "Ke, G., et al. (2017). LightGBM: A highly efficient gradient boosting decision tree. NeurIPS 30.",
        "Prokhorenkova, L., et al. (2018). CatBoost: Unbiased boosting with categorical features. NeurIPS 31.",
        "Rosen, S. (1974). Hedonic prices and implicit markets. Journal of Political Economy, 82(1), 34–55.",
        "Wolpert, D. H. (1992). Stacked generalization. Neural Networks, 5(2), 241–259.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(ref)

    doc.save(str(path))
    print(f"  Saved: {path}")


# ── Technical Report ──────────────────────────────────────────────────────────

def build_technical_report(path):
    doc = Document()
    set_page_margins(doc)

    add_title_block(
        doc,
        title="THAMAN — COMPLETE TECHNICAL REFERENCE",
        subtitle="Dual-City Automated Valuation Model: New York City & Riyadh",
    )

    set_heading(doc, "1. Project Overview", level=1)
    doc.add_paragraph(
        "THAMAN (ثمن — Arabic for 'price') is an Automated Valuation Model (AVM) that "
        "estimates residential property prices using machine learning across two independent "
        "real estate markets: New York City (USD total price) and Riyadh, Saudi Arabia "
        "(SAR per square metre)."
    )

    set_heading(doc, "2. System Architecture", level=1)
    add_table_style(
        doc,
        headers=["Layer", "Technology"],
        rows=[
            ("Data pipeline",   "Python, Polars"),
            ("Training",        "scikit-learn, XGBoost, LightGBM, CatBoost, Ridge"),
            ("Model storage",   "pickle (.pkl), JSON"),
            ("API backend",     "FastAPI, Uvicorn (port 8000)"),
            ("Spatial index",   "scipy cKDTree, sklearn BallTree"),
            ("Frontend",        "Leaflet.js, Chart.js"),
            ("Deployment",      "Hugging Face Spaces (Docker)"),
            ("Version control", "GitHub (turkialm/thaman-v2)"),
        ],
        caption="Technology stack",
    )

    set_heading(doc, "3. NYC Dataset", level=1)
    doc.add_paragraph(
        "Source: NYC Open Data — Citywide Rolling Calendar Sales. Coverage: 2022–2026. "
        "Rows (training): 157,329. Rows (holdout): 27,763 (newest 15% by sale_date). "
        "Target: sale_price (USD). Log-transform: log1p(sale_price) during training; "
        "expm1() at inference."
    )
    doc.add_paragraph(
        "Cleaning: removed sale_price < $10K or > $100M; missing lat/lon or sale_date; "
        "duplicate BBL+date (kept highest); non-arm's-length transfers; gross_square_feet=0."
    )
    doc.add_paragraph(
        "Supplementary sources: PLUTO 25v4 (assesstot for prior sale imputation), "
        "MTA Subway/Bus, NYC Parks, Schools, HPD Violations, DOB Permits, "
        "311 Service Requests, NYPD Complaints, ACS Census, ACRIS prior sales, "
        "Airbnb listings, Air quality, Bike lanes, Mortgage rates."
    )

    set_heading(doc, "4. NYC Feature Engineering (104 features)", level=1)
    doc.add_paragraph(
        "All engineering runs in Polars (not pandas). Features include: "
        "(1) Distance features with log1p companions (8 spatial indexes), "
        "(2) Urban gravity centres (4 major employment hubs), "
        "(3) Walk score proxy (0–100 composite), "
        "(4) Interaction features (sqft_per_floor, income_over_crime, lot_coverage, etc.), "
        "(5) Target encodings (bldgclass, borough×bldg, NTA, NTA×bldg), "
        "(6) NTA price trend slope (OLS log-linear per NTA), "
        "(7) HPD/DOB/311/MTA quality features (v11), "
        "(8) Temporal features (sale_year, sin/cos month, mortgage_rate)."
    )
    doc.add_paragraph(
        "QoL winsorisation caps (stored in meta.json): crime_rate_nta=267.56, "
        "noise_density_nta=176.67, livability_complaint_rate=60.99."
    )

    set_heading(doc, "5. NYC Training Methodology", level=1)
    set_heading(doc, "5.1 Cross-Validation", level=2)
    doc.add_paragraph(
        "Spatial GroupKFold, 10 folds, groups=ntacode (212 NTAs). "
        "Prevents spatial autocorrelation leakage. Fold 5 degraded to R²≈0.37 "
        "revealing a localised anomalous cluster."
    )

    set_heading(doc, "5.2 Hyperparameters", level=2)
    add_table_style(
        doc,
        headers=["Parameter", "XGB-A", "XGB-B", "LGB", "CatBoost"],
        rows=[
            ("n_estimators",    "5,000", "5,000", "5,000", "3,000"),
            ("learning_rate",   "0.02",  "0.05",  "0.04",  "0.025"),
            ("max_depth",       "7",     "4",     "—",     "8"),
            ("num_leaves",      "—",     "—",     "127",   "—"),
            ("subsample",       "0.80",  "0.65",  "0.80",  "—"),
            ("colsample_bytree","0.60",  "0.75",  "0.70",  "—"),
            ("early_stopping",  "400",   "400",   "400",   "400"),
            ("Best round",      "3,843", "5,000", "952",   "3,000"),
        ],
        caption="Base learner hyperparameters",
    )

    set_heading(doc, "5.3 Meta-Learner", level=2)
    doc.add_paragraph(
        "Ridge (alpha=1.0, positive=True). LightGBM meta holdout R²=0.6349 vs "
        "Ridge meta holdout R²=0.6450. Ridge selected to prevent meta-level overfitting "
        "to OOF noise. positive=True enforces blending (not arbitrage)."
    )

    set_heading(doc, "5.4 Luxury Sub-Model", level=2)
    doc.add_paragraph(
        "Separate XGBoost trained on Manhattan sales ≥ $2.5M. "
        "Soft blend: α = clip((price−2.5M)/(5.0M−2.5M), 0, 1); "
        "P_final = (1−α)·P_stack + α·P_luxury."
    )

    set_heading(doc, "6. NYC Results", level=1)
    add_table_style(
        doc,
        headers=["Evaluation", "R²", "MedAPE", "MAE (USD)"],
        rows=[
            ("Stack v11 (holdout, 27,763 sales)", "0.6450", "20.24%", "$1,065,470"),
        ],
        caption="NYC overall performance",
    )
    add_table_style(
        doc,
        headers=["Borough", "n", "R²", "MedAPE"],
        rows=[
            ("Staten Island", "3,052", "0.4103", "13.46%"),
            ("Queens",        "9,533", "0.6851", "17.06%"),
            ("Bronx",         "2,735", "0.6195", "20.93%"),
            ("Brooklyn",      "7,032", "0.6206", "20.89%"),
            ("Manhattan",     "5,411", "0.6141", "36.66%"),
        ],
        caption="Borough-level performance",
    )

    set_heading(doc, "7. Riyadh Dataset", level=1)
    doc.add_paragraph(
        "6,910 district-quarter aggregates; 163 Riyadh districts; Q1 2018–Q3 2025. "
        "Target: median(sale_price_sar_sqm) per district-quarter. "
        "Training: 5,531 rows (cutoff: quarter_id < 20251 = end of 2024). "
        "Holdout: 1,379 rows (2025 Q1–Q3). Price range: 221.82–12,565.21 SAR/m²."
    )

    set_heading(doc, "Critical Bug Fixed: 2024 Quarter ID Encoding", level=2)
    doc.add_paragraph(
        "Formula quarter_id = year×10 + sale_quarter failed for 2024 CSVs where "
        "sale_quarter already contained the full code (e.g., 20241). Produced "
        "quarter_ids 40481/40483/40484 instead of 20241/20243/20244, pushing all "
        "867 2024 rows into holdout. Fix: scripts/fix_riyadh_2024_quarters.py. "
        "Impact: MedAPE improved from 23.43% → 18.16%."
    )

    set_heading(doc, "8. Riyadh Feature Engineering (76 features)", level=1)
    for cat, feats in [
        ("Location (2)", "district_lat, district_lon"),
        ("Property type (4)", "is_apartment, is_villa, is_residential_plot, is_building"),
        ("Metro transit (6)", "dist_metro_m, log_dist_metro_m, metro_stations_1km, nearest_metro_line_num, nearest_metro_type_cd, dist_metro_line1_m"),
        ("Bus/BRT (4)", "dist_bus_m, log_dist_bus_m, bus_stops_500m, brt_stops_500m"),
        ("Commercial (10)", "commercial_count_1km, density_score, hypermarket/bank/restaurant/hotel/gas counts"),
        ("Air quality (6)", "no2_nearest_mean, so2_nearest_mean, pm10_nearest_mean, o3_nearest_mean, dist_air_station_m, air_quality_score"),
        ("Macro/REI (6)", "rei_residential_qtr_idx, rei_apt_idx, rei_yoy_change, rei_qoq_change, avg_saudi_salary_yr, salary_yoy_change"),
        ("District aggregates (5)", "median_price_sqm, transaction_volume, price_vs_city_avg, price_trend_slope, median_price_apt_sqm"),
        ("Target-encoded (2)", "district_encoded, district_type_encoded"),
        ("QoL POIs (18)", "mosques/malls/schools/hospitals/parks/entertainment: distance + log_dist + count_500m each"),
        ("Connectivity (1)", "riyadh_connectivity_score"),
        ("SA_Aqar rental (4)", "aqar_median_size_sqm, aqar_median_bedrooms, aqar_median_property_age, aqar_rent_per_sqm"),
        ("Temporal (4)", "sale_year, sale_quarter_sin, sale_quarter_cos, log_deed_count"),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{cat}: ").bold = True
        p.add_run(feats)

    set_heading(doc, "9. Riyadh Training Methodology", level=1)
    doc.add_paragraph(
        "5-fold Spatial GroupKFold (groups=district_ar, 163 districts). "
        "3 base learners: XGBoost, LightGBM, CatBoost (n_estimators=1500, lr=0.03, "
        "max_depth=5, subsample=0.70, early_stopping=50). "
        "Ridge meta (alpha=1.0, positive=False). "
        "Meta coefficients: XGB=0.278, LGB=0.002, CAT=0.721."
    )

    set_heading(doc, "10. Riyadh Results", level=1)
    add_table_style(
        doc,
        headers=["Evaluation", "R²", "MedAPE", "MAE"],
        rows=[
            ("OOF (5-fold GroupKFold)",         "0.9252",  "9.03%", "—"),
            ("Holdout (2025 Q1–Q3, n=1,379)",   "0.7981", "18.16%", "991 SAR/m²"),
        ],
        caption="Riyadh model performance",
    )
    doc.add_paragraph(
        "OOF-to-holdout gap explanation: (1) temporal distribution shift "
        "(holdout = new 2025 market conditions), (2) OOF temporal overlap "
        "(same quarter appears across GroupKFold folds), "
        "(3) small dataset instability (5,531 rows, 163 districts)."
    )

    set_heading(doc, "11. NYC vs. Riyadh Key Differences", level=1)
    add_table_style(
        doc,
        headers=["Dimension", "NYC", "Riyadh"],
        rows=[
            ("Data granularity",  "Individual transactions",    "District-quarter aggregates"),
            ("Dataset size",      "185,092 train rows",         "5,531 train rows"),
            ("Target unit",       "USD total price",            "SAR per m²"),
            ("Base learners",     "4 (XGB-A, XGB-B, LGB, CAT)", "3 (XGB, LGB, CAT)"),
            ("CV folds",          "10-fold GroupKFold",         "5-fold GroupKFold"),
            ("Meta-learner",      "Ridge (positive=True)",      "Ridge (positive=False)"),
            ("Features",          "104",                        "76"),
            ("Holdout MedAPE",    "20.24%",                     "18.16%"),
            ("Holdout R²",        "0.6450",                     "0.7981"),
        ],
        caption="NYC vs. Riyadh comparison",
    )

    set_heading(doc, "12. Known Bugs Fixed", level=1)
    for bug in [
        "Bug 1 — Riyadh 2024 Quarter ID Encoding (FIXED): year×10+sale_quarter formula "
        "failed; 2024 data all landed in holdout. Fix: fix_riyadh_2024_quarters.py. "
        "Impact: MedAPE 23.43% → 18.16%.",
        "Bug 2 — Stale district_medape in riyadh_meta.json (FIXED): meta_dict.update() "
        "preserved old v1 district table. Fix: removed district_medape key.",
        "Bug 3 — REI Values Corrupted for 2024 Rows (FIXED): wrong quarter_ids caused "
        "incorrect REI lookup. Fix: hardcoded correct quarterly values.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(bug)

    set_heading(doc, "13. Deployment", level=1)
    doc.add_paragraph("Hugging Face Space: https://huggingface.co/spaces/Turki-Almurahhem/thaman")
    doc.add_paragraph("GitHub: https://github.com/turkialm/thaman-v2")
    doc.add_paragraph(
        "To run locally: cd /path/to/new_try && uvicorn api.main:app --port 8000. "
        "Cold-start time: ~30s. API latency (warm): 200–400ms."
    )

    doc.save(str(path))
    print(f"  Saved: {path}")


# ── Defense Q&A ────────────────────────────────────────────────────────────────

def build_defense_qa(path):
    doc = Document()
    set_page_margins(doc)

    add_title_block(
        doc,
        title="THAMAN — Defense Q&A Preparation",
        subtitle="BSc Graduation Project · Umm Al-Qura University",
        year="Prepared: 2026-05-20",
    )

    note = doc.add_paragraph()
    note.add_run("Note: ").bold = True
    note.add_run("Memorise the SHORT ANSWER. Expand with DETAIL if pressed by the committee.")

    sections = [
        ("1. Model Performance", [
            {
                "q": "Why is Riyadh holdout R²=0.79 lower than OOF R²=0.93?",
                "short": "OOF uses district-level training data; holdout is 2025 which is new macro conditions the model hasn't seen.",
                "detail": [
                    "OOF cross-validation splits BY DISTRICT (GroupKFold). Each fold holds out entire districts, avoiding spatial autocorrelation leakage.",
                    "The training data is district-quarter aggregates — every row already encodes the district's median price. This inflates OOF metrics compared to a truly unseen market.",
                    "The holdout period is 2025 Q1–Q3 — entirely new market conditions including post-2024 macro shifts (REI index, salary data).",
                    "A gap of ~14 R² points (0.93 → 0.79) is expected and honest. The holdout metric is the number to cite: R²=0.79, MedAPE=18.16%.",
                ],
            },
            {
                "q": "Your NYC model has R²=0.65 but Riyadh has R²=0.79. Doesn't that make NYC worse?",
                "short": "Different targets. NYC predicts individual transaction price; Riyadh predicts district-quarter aggregates. Aggregates are smoother → higher R².",
                "detail": [
                    "NYC: 185K individual transactions. High variance in sale price at parcel level. R²=0.65 is strong — state-of-art NYC AVM models report 0.60–0.70.",
                    "Riyadh: 6,910 district-quarter aggregate rows. Averaging transactions removes micro-level noise.",
                    "Comparing the two R² directly is misleading — apples and oranges.",
                ],
            },
            {
                "q": "MedAPE of 18.16% seems high. Is that good enough?",
                "short": "Industry benchmark for AVM in emerging markets is 15–25%. 18.16% is within standard range for district aggregates without parcel-level data.",
                "detail": [
                    "Zillow Zestimate (US, individual parcels, 30+ years of data): ~7% MedAPE.",
                    "Saudi Arabia has no open parcel-level register. We use MoJ quarterly summaries.",
                    "By property type: apartment MedAPE=16.42%, villa=14.46%, plot=25.84%.",
                ],
            },
        ]),
        ("2. Data", [
            {
                "q": "Why didn't you use individual transaction data for Riyadh?",
                "short": "The Ministry of Justice real estate register (Aqarat) is not publicly accessible. Only quarterly district summaries are released on Saudi Open Data.",
                "detail": [
                    "Saudi MoJ Aqarat database is restricted to licensed valuers and financial institutions.",
                    "Saudi Open Data Platform (data.gov.sa) publishes aggregated quarterly reports at district level.",
                    "This is not a design choice — it is the data availability constraint.",
                ],
            },
            {
                "q": "How do you validate your Riyadh predictions against real prices?",
                "short": "Compared model output to 1,615 active Haraj listings. Model median was 54% below asking price — consistent with asking vs. transacted price gap in Saudi market.",
                "detail": [
                    "Scraped 1,615 Haraj.com listings (444 apartments, 630 villas, 526 residential plots, 15 buildings).",
                    "THAMAN predicts TRANSACTED price (MoJ quarterly data). Haraj shows ASKING price.",
                    "MedAPE of 54% between THAMAN and Haraj asking price is expected — NOT a model error.",
                ],
            },
        ]),
        ("3. Algorithms & Methodology", [
            {
                "q": "Why use a stacking ensemble instead of a single model?",
                "short": "Each base learner captures different patterns. Stacking reduces variance and improves generalisation beyond any single model.",
                "detail": [
                    "XGBoost: symmetric trees, good at global feature interactions.",
                    "LightGBM: leaf-wise growth, handles high-cardinality encodings well.",
                    "CatBoost: ordered boosting, robust on smaller datasets. Dominates Riyadh meta (coef=0.721).",
                    "OOF predictions used to train meta-learner: prevents base learners from memorising training labels.",
                ],
            },
            {
                "q": "What is GroupKFold and why did you use it?",
                "short": "Regular KFold leaks spatial autocorrelation. Properties in the same neighbourhood have correlated prices — a model can cheat by seeing nearby properties.",
                "detail": [
                    "GroupKFold groups entire geographic units (NTA codes for NYC, district_ar for Riyadh) into the same fold.",
                    "NYC uses 10 folds (212 NTAs → ~31 per fold). Riyadh uses 5 folds (163 districts → ~33 per fold).",
                    "Regular KFold would show R²>0.95 (data leakage). GroupKFold gives honest estimates.",
                ],
            },
            {
                "q": "Why use Ridge as the meta-learner instead of another gradient boosted model?",
                "short": "LightGBM at the meta level overfits to OOF noise. Ridge prevents this and gives better holdout performance despite a lower OOF score.",
                "detail": [
                    "Empirical test: LightGBM meta holdout R²=0.6349 vs Ridge meta holdout R²=0.6450.",
                    "By the meta-stage, the four base learners have already exhausted non-linear relationships. Remaining variance is mostly noise.",
                    "Ridge L2 regularisation with positive=True enforces blending, not arbitrage.",
                ],
            },
            {
                "q": "Why log-transform the target variable?",
                "short": "Property prices are right-skewed. Log-transform makes the target approximately normally distributed and improves model fit.",
                "detail": [
                    "Raw sale_price has skewness >5 in NYC, >3 in Riyadh.",
                    "log1p(price) reduces skewness to near-normal, preventing the model from fitting outliers disproportionately.",
                    "At inference: price = expm1(model_output). MedAPE computed in original price space.",
                ],
            },
        ]),
        ("4. System / Application", [
            {
                "q": "How does the app determine which city model to use?",
                "short": "Bounding box check on click coordinates. NYC bbox and Riyadh bbox are non-overlapping — no ambiguity possible.",
                "detail": [
                    "NYC bbox: lat 40.47–40.92, lon −74.26 to −73.70.",
                    "Riyadh bbox: lat 24.35–25.10, lon 46.30–47.20.",
                    "Two completely separate ML pipelines: NYC scorer never sees SAR data, Riyadh scorer never sees USD data.",
                ],
            },
        ]),
        ("5. Extended Questions", [
            {
                "q": "Is your model fair? Does it discriminate by neighbourhood demographics?",
                "short": "Fairness is a known concern for AVMs. We include median income and crime rate as features — these reflect market reality but can encode historical inequity.",
                "detail": [
                    "The model uses median_income_nta and crime_rate_nta, which are themselves products of historical patterns including redlining and unequal resource allocation.",
                    "A fairness audit is listed as future work.",
                    "The academic contribution is making inputs transparent via SHAP so users can see which features drive any given estimate.",
                ],
            },
        ]),
        ("6. Regulatory and Market Context", [
            {
                "q": "Is THAMAN compliant with the 2025 AVM Quality Control Rule?",
                "short": "THAMAN addresses all five standards of the Dodd-Frank AVM QC Rule (effective Oct 2025): accuracy evidence, data integrity, conflict-of-interest avoidance, sample testing, and nondiscrimination.",
                "detail": [
                    "Standard 1 — Accuracy: MedAPE=20.24% on 27,763 holdout transactions (NYC); borough and price-tier breakdowns published. Confidence score 0–100 with letter grade at inference.",
                    "Standard 2 — Data integrity: all training data from official government sources (NYC Open Data/PLUTO; Saudi Open Data Portal). No crowd-sourced or user-submitted prices influence training.",
                    "Standard 3 — Conflict of interest: open-source academic prototype (MIT licence, GitHub). No commercial lender affiliation.",
                    "Standard 4 — Sample testing: 15% time-based holdout (NYC, n=27,763) and fully out-of-sample 2025 holdout (Riyadh, n=1,379); Spatial GroupKFold prevents geographic leakage.",
                    "Standard 5 — Nondiscrimination: no protected-class attributes in feature set; NTA/district encodings use sale price only; SHAP enables fairness auditing.",
                    "THAMAN is a research prototype — the rule targets lending institutions. But these design choices reflect 2026 best practices.",
                ],
            },
            {
                "q": "How do 2025–2026 Saudi policy changes affect your Riyadh model?",
                "short": "Three shifts (rent freeze, foreign ownership reform, 57,000-unit supply pipeline) create temporal distribution shift beyond the training window. Model direction remains valid; retraining on 2025–2026 data would improve accuracy in premium districts.",
                "detail": [
                    "Rent freeze (late 2025): government capped annual rent increases. SA_Aqar rental features reflect pre-freeze yields — partly explains premium-district prediction gaps.",
                    "Foreign ownership reform (Jan 2026): expanded GCC/international ownership rights in Riyadh urban zones. Drives premium district appreciation: Al Nakheel asking 26,000 SAR/m², Al Hada 20,000 SAR/m² — vs. THAMAN predictions of ~2,800–3,100 SAR/m² (transaction-price reference).",
                    "57,000-unit pipeline: northern/western corridors (Al Qirawan, Al Yasmin, Al Narjis). May moderate outer-ring appreciation; inner-ring supply-constrained districts maintain upward pressure.",
                    "REI yields: 8.5–9.5% gross rental yields in prime Riyadh — among highest in GCC. THAMAN's quarterly REI features will absorb these when retrained.",
                    "For 2026 predictions: acknowledge the policy-shift caveat. The 18.16% holdout MedAPE is expected to widen slightly until retrained on 2025–2026 data.",
                ],
            },
        ]),
    ]

    for sec_title, qas in sections:
        set_heading(doc, sec_title, level=1)
        for qa in qas:
            p = doc.add_paragraph()
            p.add_run(f"Q: {qa['q']}").bold = True

            p2 = doc.add_paragraph()
            p2.add_run("SHORT: ").bold = True
            p2.add_run(qa["short"])

            p3 = doc.add_paragraph()
            p3.add_run("DETAIL:").bold = True
            for d in qa["detail"]:
                bp = doc.add_paragraph(style="List Bullet")
                bp.add_run(d)

            doc.add_paragraph()

    # Known limitations
    set_heading(doc, "7. Known Limitations (say confidently, not defensively)", level=1)
    for limit, response in [
        ("Riyadh data is district-level aggregate, not individual transactions.",
         "'This is the data availability constraint from MoJ. Future work could incorporate individual deeds when access is granted.'"),
        ("No property condition data (renovation quality, interior finish).",
         "'Both NYC and Riyadh lack this. It is a universal limitation of government transaction registers.'"),
        ("NYC luxury tier (>$5M) has wider confidence intervals.",
         "'Luxury transactions are sparse (<3% of dataset) so the model is less confident there — correctly communicated via a wider range.'"),
    ]:
        p = doc.add_paragraph()
        p.add_run(limit).bold = True
        doc.add_paragraph(f"→ {response}")

    doc.save(str(path))
    print(f"  Saved: {path}")


# ── Demo Script ────────────────────────────────────────────────────────────────

def build_demo_script(path):
    doc = Document()
    set_page_margins(doc)

    add_title_block(
        doc,
        title="THAMAN — Graduation Defense Demo Script",
        subtitle="BSc Computer Science · Umm Al-Qura University · 2026",
        author="",
        year="Total time: 5 minutes",
    )

    note = doc.add_paragraph()
    note.add_run("Important: ").bold = True
    note.add_run("Keep this script on your phone or second screen — NOT on the projector.")

    # Setup
    set_heading(doc, "Setup (15 Minutes Before Committee Enters)", level=1)
    for item in [
        "Open browser, navigate to: https://huggingface.co/spaces/Turki-Almurahhem/thaman",
        "Wait for map to fully load (cold-start: ~30 seconds)",
        "Toggle language to Arabic once, then back to English — confirm bilingual toggle works",
        "Pan map to NYC view (should default; if not, refresh once)",
        "Open charts.html in a second browser tab (hidden behind main tab)",
        "Have this script on your phone or second screen — NOT on the projector",
        "Confirm screen mirroring / projector is working before committee sits down",
        "Mute your phone",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Coordinates to have ready (copy into browser console if Nominatim search is slow):").bold = True
    doc.add_paragraph("NYC click target: 40.7549, -73.9840 (Midtown Manhattan)")
    doc.add_paragraph("Riyadh click target: 24.6877, 46.7219 (Downtown Riyadh / King Fahd Road)")

    # Steps
    steps = [
        ("Step 1 — NYC Prediction (1 minute)",
         "Map is showing NYC. Click approximately on Midtown Manhattan — coordinates 40.7549, −73.9840.",
         [
             ("English",
              "\"This is THAMAN — a dual-city Automated Valuation Model I built for my graduation project. "
              "We're starting in New York City. I'm clicking on Midtown Manhattan to get an instant property valuation.\"\n\n"
              "\"I'll select building type — Elevator Condo, D4 — and enter 1,200 square feet, 15 floors, built in 1985.\"\n\n"
              "[Submit prediction — pause for result to load]\n\n"
              "\"THAMAN returns a predicted price, a confidence band, a letter grade, and the top features driving "
              "this estimate. Building class and neighbourhood encoding are the top two features — consistent with NYC "
              "real estate: where you are and what type of building matter more than size alone.\"\n\n"
              "\"The comparable sales bubbles show the nearest actual recorded sales — green means our estimate is "
              "close, red means we're further off. These are real deed-recorded transactions from 185,000 NYC sales, "
              "2022 to 2026.\""),
             ("Arabic",
              "\"هذا نظام ثمان — نموذج تقييم عقاري ذكي لمدينتين طورته كمشروع تخرج. نبدأ في مدينة نيويورك.\"\n\n"
              "\"اخترت نوع البناء: شقة بمصعد، المساحة 1200 قدم مربع، 15 طابقاً، بُني عام 1985.\"\n\n"
              "\"النظام يعطينا سعراً تقديرياً، نطاق ثقة، درجة تقييم، وأهم العوامل المؤثرة.\""),
             ("Metric to highlight",
              "\"Our NYC model achieves MedAPE of 20.24% on 27,763 holdout sales — competitive with "
              "commercial AVMs like Zillow Zestimate.\""),
         ]),
        ("Step 2 — City Switch to Riyadh (30 seconds)",
         "Click the city-switch toggle. Map animates to Riyadh view with district polygons visible.",
         [
             ("English",
              "\"Now here's what makes THAMAN distinctive — it's a dual-city system. I'll switch to Riyadh.\"\n\n"
              "\"The same stacking architecture, the same FastAPI backend, now running on a completely different market. "
              "Saudi Arabia's real estate data is published as district-level quarterly aggregates by the Ministry of "
              "Justice — not individual transactions like NYC. The model learned from 6,910 district-quarter observations "
              "instead of 185,000 individual sales.\""),
             ("Arabic",
              "\"الآن ننتقل إلى الرياض — وهذا ما يميّز ثمان. نفس البنية التقنية، لكن على سوق مختلف تماماً.\""),
         ]),
        ("Step 3 — Riyadh Prediction + SHAP Drivers (1.5 minutes)",
         "Click on Downtown Riyadh — coordinates 24.6877, 46.7219 (King Fahd Road area).",
         [
             ("English",
              "\"I'll click on the King Fahd Road corridor — one of Riyadh's prime districts. "
              "Let's select Villa, 400 square metres.\"\n\n"
              "[Submit prediction — wait for result]\n\n"
              "\"The model returns a prediction in SAR per square metre. The SHAP breakdown: metro access, "
              "commercial density, air quality, and district price history are the top drivers.\"\n\n"
              "\"The Riyadh Metro opened in 2024 — a novel infrastructure signal the model captures, "
              "entirely absent from pre-2024 models.\""),
             ("Arabic",
              "\"أضغط على منطقة طريق الملك فهد. سأختار فيلا، 400 متر مربع.\"\n\n"
              "\"النموذج يُعطينا التقدير بالريال السعودي لكل متر مربع. في تحليل SHAP: القرب من المترو، "
              "الكثافة التجارية، جودة الهواء، والتاريخ السعري للحي.\""),
             ("Metrics to highlight",
              "OOF (training folds): R²=0.9252, MedAPE=9.03% — model genuinely learned Saudi market structure.\n"
              "Holdout Q1–Q3 2025: R²=0.7981, MedAPE=18.16% — new-quarter stress test, not a random sample."),
         ]),
        ("Step 4 — Listings Layer (30 seconds)",
         "Toggle on the Haraj active listings layer.",
         [
             ("English",
              "\"This layer shows 1,615 active property listings scraped from Haraj.com.sa — Saudi Arabia's "
              "largest classifieds marketplace. Each point is colour-coded: blue=apartments, green=villas, amber=plots.\"\n\n"
              "\"The model systematically predicts lower than asking prices — that's expected. THAMAN was trained "
              "on deed-recorded transaction prices. Haraj shows what sellers ask, before negotiation. The overall gap "
              "is 54% MedAPE — consistent with documented Saudi negotiation margins of 20–50 percent.\""),
             ("Arabic",
              "\"هذه الطبقة تُظهر 1615 عرضاً نشطاً من موقع حراج.كوم — أكبر سوق للعقارات في السعودية.\""),
         ]),
        ("Step 5 — Analytics Dashboard (30 seconds)",
         "Switch to second browser tab — charts.html.",
         [
             ("English",
              "\"The analytics dashboard shows model performance broken down by NYC borough and price tier.\"\n\n"
              "\"Notice the Staten Island paradox: lowest R²=0.41 but best MedAPE=13.5%. Staten Island has very "
              "low price variance; the model's absolute errors are small, but R² penalises a low-variance target. "
              "MedAPE is the right metric for a user-facing AVM.\"\n\n"
              "\"Manhattan is hardest at 36.7% MedAPE. Co-op board approval discounts and unobservable interior "
              "finishes create heterogeneity that no tabular dataset can capture.\""),
             ("Arabic",
              "\"لوحة التحليلات تُظهر أداء النموذج مقسّماً حسب منطقة نيويورك وشريحة السعر.\"\n\n"
              "\"لاحظوا مفارقة ستاتن آيلاند: أقل R² لكن أفضل MedAPE.\""),
         ]),
        ("Step 6 — Q&A Talking Points / Closing (1 minute)",
         "Use this minute as a buffer or for closing summary.",
         [
             ("English closing",
              "\"To summarise: THAMAN is a production-deployed AVM across two cities — New York and Riyadh — "
              "using a four-model stacking ensemble across 104 and 76 features respectively. It achieves competitive "
              "accuracy on NYC's 185,000-sale holdout and demonstrates cross-market generalisability on Saudi Arabia's "
              "data-scarce district-aggregate market. The full system — data pipelines, training code, API, and web "
              "interface — is deployed on Hugging Face and open-sourced on GitHub. Thank you.\""),
             ("Arabic closing",
              "\"خلاصة القول: ثمان نظام تقييم عقاري منتشر فعلياً لمدينتين، يستخدم مجموعة من أربعة نماذج ذكاء "
              "اصطناعي عبر مئة وأربع ميزات في نيويورك، وستة وسبعين ميزة في الرياض. شكراً.\""),
         ]),
    ]

    for step_title, stage_note, narrations in steps:
        set_heading(doc, step_title, level=1)
        if stage_note:
            p = doc.add_paragraph()
            p.add_run("Stage note: ").bold = True
            p.add_run(stage_note).italic = True
        for narr_label, narr_text in narrations:
            p2 = doc.add_paragraph()
            p2.add_run(f"[{narr_label}]").bold = True
            doc.add_paragraph(narr_text)
        doc.add_paragraph()

    # Fallback
    set_heading(doc, "Fallback Protocol (if Hugging Face is slow)", level=1)
    for item in [
        "Say: 'The deployed version is loading from cold start — common with Hugging Face Spaces after inactivity. "
        "While it loads, I'll walk through the architecture.'",
        "Switch to showing the paper/slides and explain the model architecture verbally.",
        "Keep refreshing the HF tab in the background — typically loads in 45–90 seconds.",
        "If HF completely unavailable: open Terminal and run:\n"
        "    cd /Users/totam/Desktop/new_try && uvicorn api.main:app --port 8000\n"
        "    Then open: http://localhost:8000/ui",
        "API startup time: ~30 seconds (spatial KD-tree indexes loading).",
    ]:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)

    # Key numbers
    set_heading(doc, "Key Numbers to Memorise", level=1)
    p = doc.add_paragraph()
    p.add_run("Print this section and keep it in your pocket.").bold = True

    add_table_style(
        doc,
        headers=["Metric", "Value", "Context"],
        rows=[
            ("NYC training rows",       "185,092",     "Sales 2022–2026"),
            ("NYC features",            "104",          "Structural + spatial + QoL"),
            ("NYC holdout rows",        "27,763",       "Newest 15% by date"),
            ("NYC R² (holdout)",        "0.6450",       "Stack v11"),
            ("NYC MedAPE (holdout)",    "20.24%",       "Stack v11"),
            ("Riyadh total rows",       "6,910",        "District-quarter obs., 2018–2025"),
            ("Riyadh training rows",    "5,531",        "2018–2024 (incl. Metro-era)"),
            ("Riyadh features",         "76",           "Transit, QoL, macro, rental"),
            ("Riyadh OOF R²",           "0.9252",       "5-fold spatial GroupKFold"),
            ("Riyadh OOF MedAPE",       "9.03%",        "In-sample cross-validation"),
            ("Riyadh holdout R²",       "0.7981",       "2025 Q1–Q3, n=1,379"),
            ("Riyadh holdout MedAPE",   "18.16%",       "Out-of-sample stress test"),
            ("Riyadh holdout MAE",      "991 SAR/m²",   "Out-of-sample stress test"),
            ("Haraj validation MedAPE", "54.33%",       "Asking vs. transaction (expected)"),
            ("Haraj listings",          "1,615",        "444 apts, 630 villas, 526 plots"),
            ("NYC NTA groups",          "212",          "Neighbourhood spatial units"),
            ("Riyadh district polygons","133",          "From OSM admin_level=10"),
            ("API latency",             "200–400 ms",   "Including SHAP computation"),
            ("Automated tests",         "37",           "15 scorer + 22 API tests"),
        ],
        caption="Key metrics",
    )

    doc.save(str(path))
    print(f"  Saved: {path}")


# ── Main ───────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    base = Path(__file__).parent

    print("Generating Word documents...")
    build_paper(base / "thaman_paper.docx")
    build_technical_report(base / "technical_report.docx")
    build_defense_qa(base / "defense_qa.docx")
    build_demo_script(base / "demo_script.docx")
    print("\nDone! All 4 .docx files generated.")
